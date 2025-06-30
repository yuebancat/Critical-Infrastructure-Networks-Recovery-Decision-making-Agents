import os
import pandas as pd
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
import json

# RAG Tool: Retrieve chunks for each question
def RAG(input_docx):
    """
    This tool retrieves the top 3 relevant chunks for each question from a docx file using RAG.

    :param input_docx: Path to the .docx file containing tool descriptions.
    :param input_questions_excel: Path to the Excel file containing questions.
    :return: DataFrame with questions and the corresponding 3 retrieved chunks.
    """
    # Load the tool descriptions from the docx file
    input_docx='tool_chunks_version_R2.docx'
    input_questions_excel='question4UI.xlsx'
    # Set up environment variables
    os.environ["http_proxy"] = "http://localhost:7890"
    os.environ["https_proxy"] = "http://localhost:7890"
    os.environ["OPENAI_API_KEY"] = "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"

    def extract_text_from_docx(file_path):
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

    text = extract_text_from_docx(input_docx)
    docs = [Document(page_content=text, metadata={"source": input_docx})]

    # Split the document into smaller chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50, add_start_index=True)
    all_splits = text_splitter.split_documents(docs)

    # Load the questions from the Excel file
    questions_df = pd.read_excel(input_questions_excel)

    # Set up embeddings and vector store
    embedding_model = OpenAIEmbeddings(
        model="text-embedding-3-small",
        openai_api_key=os.environ["OPENAI_API_KEY"]
    )
    vector_store = FAISS.from_documents(all_splits, embedding_model)

    # Prepare the output DataFrame
    output_df = pd.DataFrame(columns=["question", "retrieved_chunk_1", "retrieved_chunk_2", "retrieved_chunk_3"])

    # Retrieve chunks for each question
    for index, row in questions_df.iterrows():
        user_question = row["questions"]

        # Retrieve relevant documents using similarity search
        retrieved_docs = vector_store.similarity_search(user_question, k=3)

        # Extract the context from retrieved chunks
        context = "\n\n".join([doc.page_content for doc in retrieved_docs])

        # Add results to the output dataframe
        output_df = output_df.append({
            "question": user_question,
            "retrieved_chunk_1": retrieved_docs[0].page_content,
            "retrieved_chunk_2": retrieved_docs[1].page_content,
            "retrieved_chunk_3": retrieved_docs[2].page_content
        }, ignore_index=True)

    # Save the output DataFrame to Excel
    output_file_path = "question_with_retrieved_chunks.xlsx"
    output_df.to_excel(output_file_path, index=False)

    print(f"RAG results saved to {output_file_path}")
    return output_file_path

